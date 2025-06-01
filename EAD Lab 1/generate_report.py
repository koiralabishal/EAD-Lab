from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm

def add_horizontal_line(paragraph):
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    pPr.insert_element_before(pBdr,
        'w:shd', 'w:tabs', 'w:suppressAutoHyphens', 'w:kinsoku', 'w:wordWrap',
        'w:overflowPunct', 'w:topLinePunct', 'w:autoSpaceDE', 'w:autoSpaceDN',
        'w:bidi', 'w:adjustRightInd', 'w:snapToGrid', 'w:spacing', 'w:ind',
        'w:contextualSpacing', 'w:mirrorIndents', 'w:suppressOverlap', 'w:jc',
        'w:textDirection', 'w:textAlignment', 'w:textboxTightWrap',
        'w:outlineLvl', 'w:divId', 'w:cnfStyle', 'w:rPr', 'w:sectPr',
        'w:pPrChange'
    )
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'auto')
    pBdr.append(bottom)

def add_command_with_output(doc, command, output=None):
    # Add command
    p = doc.add_paragraph()
    p.add_run('$ ').font.color.rgb = RGBColor(147, 197, 253)  # Light blue color for prompt
    p.add_run(command).font.name = 'Consolas'
    
    # Add output if provided
    if output:
        p = doc.add_paragraph()
        for line in output:
            p.add_run(line + '\n').font.name = 'Consolas'

def main():
    # Create document
    doc = Document()
    
    # Set margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(1.27)  # 0.5 inches
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)
    
    # Add college logo
    logo_paragraph = doc.add_paragraph()
    logo_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    logo_run = logo_paragraph.add_run()
    logo_run.add_picture('static/Screenshot 2025-06-01 125244.png', width=Inches(2))
    
    # Add college name
    college_name = doc.add_paragraph()
    college_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    college_run = college_name.add_run('Gandaki College of Engineering and Science')
    college_run.bold = True
    college_run.font.size = Pt(16)
    
    # Add title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run('Enterprise Application Development')
    title_run.bold = True
    title_run.font.size = Pt(14)
    
    # Add lab title
    lab_title = doc.add_paragraph()
    lab_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    lab_run = lab_title.add_run('Lab Report: Git and GitHub Demonstration')
    lab_run.italic = True
    lab_run.font.size = Pt(12)
    
    # Add lab number
    lab_num = doc.add_paragraph()
    lab_num.alignment = WD_ALIGN_PARAGRAPH.CENTER
    lab_num.add_run('Lab: 01')
    
    # Add some space
    doc.add_paragraph()
    
    # Create a table for submission details with center alignment
    table = doc.add_table(rows=2, cols=2)
    table.style = 'Table Grid'
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Make cells equal width
    for cell in table.columns[0].cells:
        cell.width = Inches(3)
    for cell in table.columns[1].cells:
        cell.width = Inches(3)
    
    # Submitted By (Left cell)
    cell_left = table.cell(0, 0)
    cell_left.text = "Submitted By:"
    p = cell_left.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].bold = True
    p = cell_left.add_paragraph("Bishal Koirala\nRoll No: 15")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Submitted To (Right cell)
    cell_right = table.cell(0, 1)
    cell_right.text = "Submitted To:"
    p = cell_right.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].bold = True
    p = cell_right.add_paragraph("Er. Prativa Nyaupane\nLecturer, Gandaki College of Engineering and Science")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # GitHub Details (Center cell spanning both columns)
    github_cell = table.cell(1, 0)
    table.cell(1, 1).merge(github_cell)
    github_details = github_cell.add_paragraph()
    github_details.alignment = WD_ALIGN_PARAGRAPH.CENTER
    github_details.add_run('GitHub Profile Name: ').bold = True
    github_details.add_run('koiralabishal\n')
    github_details.add_run('GitHub Repository Name: ').bold = True
    github_details.add_run('EAD-Lab\n')
    github_details.add_run('Repository Link: ').bold = True
    github_details.add_run('https://github.com/koiralabishal/EAD-Lab')
    
    # Add page break after cover page
    doc.add_page_break()
    
    # Add Objectives section
    doc.add_heading('Objectives', level=1)
    objectives = [
        ('1. Version Control Understanding', [
            'Master the fundamental concepts of Git version control system',
            'Understand distributed version control principles',
            'Learn repository management and initialization'
        ]),
        ('2. Branch Management', [
            'Create and manage multiple branches effectively',
            'Implement proper branching strategies',
            'Handle branch synchronization with remote repository'
        ]),
        ('3. Collaboration Skills', [
            'Master pull request workflow',
            'Practice code review processes',
            'Learn team collaboration techniques'
        ])
    ]
    
    for title, points in objectives:
        p = doc.add_paragraph()
        p.add_run(title).bold = True
        for point in points:
            doc.add_paragraph(point, style='List Bullet')
    
    # Add Theory section
    doc.add_heading('Theory', level=1)
    theory = [
        ('1. Git Version Control System', [
            'Distributed version control system for tracking changes',
            'Enables multiple developers to work simultaneously',
            'Maintains complete history of source code changes'
        ]),
        ('2. Git Repository Structure', [
            'Working Directory: Where files are edited',
            'Staging Area: Preparation for commits',
            'Local Repository: Stores commit history',
            'Remote Repository: Centralized storage on GitHub'
        ]),
        ('3. Branching and Merging', [
            'Branches allow parallel development',
            'Feature branches isolate new development',
            'Merging combines changes from different branches',
            'Pull requests facilitate code review'
        ])
    ]
    
    for title, points in theory:
        p = doc.add_paragraph()
        p.add_run(title).bold = True
        for point in points:
            doc.add_paragraph(point, style='List Bullet')
    
    # Add Implementation Process section
    doc.add_heading('Implementation Process', level=1)
    doc.add_paragraph('Step-by-step guide through the Git workflow implementation with detailed commands and explanations.')
    
    # Step 1: Project Setup
    doc.add_heading('Step 1: Project Setup', level=2)
    add_command_with_output(doc, 'mkdir EAD-Lab-01')
    add_command_with_output(doc, 'cd EAD-Lab-01')
    
    # Step 2: Git Repository Setup
    doc.add_heading('Step 2: Git Repository Setup', level=2)
    add_command_with_output(doc, 'git init', [
        'Initialized empty Git repository in D:/EAD Lab/EAD Lab 1/.git/'
    ])
    
    # Step 3: Remote Repository Integration
    doc.add_heading('Step 3: Remote Repository Integration', level=2)
    add_command_with_output(doc, 'git remote add EAD https://github.com/koiralabishal/EAD-Lab.git')
    add_command_with_output(doc, 'git remote -v', [
        'EAD     https://github.com/koiralabishal/EAD-Lab.git (fetch)',
        'EAD     https://github.com/koiralabishal/EAD-Lab.git (push)'
    ])
    
    # Step 4: Branch Management
    doc.add_heading('Step 4: Branch Management', level=2)
    add_command_with_output(doc, 'git branch -a', [
        '  bugfid',
        '* main',
        '  remotes/EAD/bugfid',
        '  remotes/EAD/develop',
        '  remotes/EAD/feature',
        '  remotes/EAD/main'
    ])
    
    # Step 5: Feature Development and Pull Request
    doc.add_heading('Step 5: Feature Development and Pull Request', level=2)
    
    # Feature Branch Work
    doc.add_heading('1. Feature Branch Work', level=3)
    add_command_with_output(doc, 'git checkout feature', [
        'Switched to branch \'feature\'',
        'Your branch is up to date with \'EAD/feature\''
    ])
    
    # Make Changes and Commit
    doc.add_heading('2. Make Changes and Commit', level=3)
    add_command_with_output(doc, 'git add index.html')
    add_command_with_output(doc, 'git commit -m "feat: Add modified EAD Lab 1"', [
        '[feature 646eaee] Add modified EAD Lab 1',
        '1 file changed, 150 insertions(+), 0 deletions(-)'
    ])
    
    # Push Feature Branch
    doc.add_heading('3. Push Feature Branch', level=3)
    add_command_with_output(doc, 'git push EAD feature', [
        'To https://github.com/koiralabishal/EAD-Lab.git',
        '   0829ef0..646eaee  feature -> feature'
    ])
    
    # Create Pull Request
    doc.add_heading('4. Create Pull Request (Feature to Develop)', level=3)
    add_command_with_output(doc, 'gh pr create --base develop --head feature --title "Add EAD Lab 1 Modified" --body "Added lab report with implementation details"', [
        'Creating pull request for feature into develop in koiralabishal/EAD-Lab',
        'https://github.com/koiralabishal/EAD-Lab/pull/1'
    ])
    
    # Step 6: Merge Process
    doc.add_heading('Step 6: Merge Process', level=2)
    
    # Merge Feature into Develop
    doc.add_heading('1. Merge Feature into Develop', level=3)
    add_command_with_output(doc, 'git checkout develop')
    add_command_with_output(doc, 'git pull EAD develop')
    add_command_with_output(doc, 'git merge feature --no-ff', [
        'Switched to branch \'develop\'',
        'Already up to date.',
        'Merge made by the \'recursive\' strategy.',
        'index.html | 150 ++++++++++++++++++',
        '1 file changed, 150 insertions(+)'
    ])
    
    # Push Develop Branch
    doc.add_heading('2. Push Develop Branch', level=3)
    add_command_with_output(doc, 'git push EAD develop', [
        'To https://github.com/koiralabishal/EAD-Lab.git',
        '   646eaee..d580c75  develop -> develop'
    ])
    
    # Create Pull Request (Develop to Main)
    doc.add_heading('3. Create Pull Request (Develop to Main)', level=3)
    add_command_with_output(doc, 'gh pr create --base main --head develop --title "Merge Develop into Main" --body "Merging completed feature work into main branch"', [
        'Creating pull request for develop into main in koiralabishal/EAD-Lab',
        'https://github.com/koiralabishal/EAD-Lab/pull/2'
    ])
    
    # Merge into Main
    doc.add_heading('4. Merge into Main', level=3)
    add_command_with_output(doc, 'git checkout main')
    add_command_with_output(doc, 'git pull EAD main')
    add_command_with_output(doc, 'git merge develop --no-ff')
    add_command_with_output(doc, 'git push EAD main', [
        'Switched to branch \'main\'',
        'Already up to date.',
        'Merge made by the \'recursive\' strategy.',
        'To https://github.com/koiralabishal/EAD-Lab.git',
        '   d580c75..e892f41  main -> main'
    ])
    
    # Clean Up Feature Branch
    doc.add_heading('5. Clean Up Feature Branch', level=3)
    add_command_with_output(doc, 'git branch -d feature')
    add_command_with_output(doc, 'git push EAD --delete feature', [
        'Deleted branch feature (was 646eaee).',
        'To https://github.com/koiralabishal/EAD-Lab.git',
        ' - [deleted]         feature'
    ])
    
    # Step 7: Final Verification
    doc.add_heading('Step 7: Final Verification', level=2)
    add_command_with_output(doc, 'git status', [
        'On branch main',
        'Your branch is up to date with \'EAD/main\'',
        'nothing to commit, working tree clean'
    ])
    
    # Add Implementation Summary
    doc.add_heading('Implementation Overview', level=2)
    summary_table = doc.add_table(rows=2, cols=4)
    summary_table.style = 'Table Grid'
    
    # Headers
    headers = ['Total Steps', 'Branches Created', 'Git Commands', 'Pull Requests']
    values = ['7', '4', '18', '2']
    
    for i, header in enumerate(headers):
        cell = summary_table.cell(0, i)
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    for i, value in enumerate(values):
        cell = summary_table.cell(1, i)
        cell.text = value
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add Conclusion section
    doc.add_heading('Conclusion', level=1)
    conclusions = [
        ('Achievements', [
            'Successfully implemented Git version control system',
            'Mastered branch creation and management',
            'Learned pull request workflow and collaboration',
            'Gained practical experience with GitHub'
        ]),
        ('Learning Outcomes', [
            'Understanding of modern version control practices',
            'Proficiency in Git commands and operations',
            'Experience with collaborative development workflows',
            'Knowledge of professional Git branching strategies'
        ])
    ]
    
    for title, points in conclusions:
        p = doc.add_paragraph()
        p.add_run(title).bold = True
        for point in points:
            doc.add_paragraph(point, style='List Bullet')
    
    # Save the document
    doc.save('EAD_Lab_Report_1.docx')

if __name__ == '__main__':
    main() 